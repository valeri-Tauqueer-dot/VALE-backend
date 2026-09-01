from __future__ import annotations

import asyncio
import io
import logging
from typing import Dict, List

from .models import MediaResult, MediaType
from .validators import DEFAULT_LIMITS, MediaValidationError

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """PDF, TXT and DOCX extraction with bounded resource usage."""

    def __init__(self) -> None:
        self.max_pages = DEFAULT_LIMITS.max_pdf_pages
        self.max_text_chars = DEFAULT_LIMITS.max_extracted_text_chars

    async def process(
        self, data: bytes, filename: str, content_type: str, user_message: str = ""
    ) -> MediaResult:
        return await asyncio.to_thread(
            self._process_sync, data, filename, content_type, user_message
        )

    def _process_sync(
        self, data: bytes, filename: str, content_type: str, user_message: str
    ) -> MediaResult:
        result = MediaResult(
            media_type=MediaType.DOCUMENT,
            filename=filename,
            content_type=content_type,
        )
        try:
            if content_type == "application/pdf":
                text, metadata = self._process_pdf(data)
            elif content_type == "text/plain":
                text, metadata = self._process_text(data)
            elif content_type == (
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ):
                text, metadata = self._process_docx(data)
            else:
                raise MediaValidationError(
                    f"Unsupported document type: {content_type}"
                )

            text = self._limit_text(text, result)
            result.extracted_text = text
            result.metadata.update(metadata)
            result.structured_data.update(
                {
                    "characters_extracted": len(text),
                    "user_question": user_message or None,
                }
            )
            result.description = "Document content was extracted successfully."
            if not text.strip():
                result.add_warning(
                    "No readable text was extracted. Scanned PDFs may require OCR."
                )
            return result
        except Exception as exc:
            logger.exception("Document processing failed")
            result.add_error(f"Document processing failed: {exc}")
            return result

    def _process_pdf(self, data: bytes) -> tuple[str, Dict]:
        try:
            import pypdf
        except ImportError as exc:
            raise RuntimeError("pypdf is not installed.") from exc

        reader = pypdf.PdfReader(io.BytesIO(data))
        if reader.is_encrypted:
            raise MediaValidationError(
                "Password-protected PDFs are not currently supported."
            )
        page_count = len(reader.pages)
        if page_count > self.max_pages:
            raise MediaValidationError(
                f"PDF has {page_count} pages. Maximum allowed is {self.max_pages}."
            )

        pages: List[str] = []
        for index, page in enumerate(reader.pages, start=1):
            try:
                page_text = page.extract_text() or ""
                if page_text.strip():
                    pages.append(f"--- PAGE {index} ---\n{page_text.strip()}")
            except Exception as exc:
                logger.warning("PDF page %s extraction failed: %s", index, exc)

        return "\n\n".join(pages).strip(), {
            "page_count": page_count,
            "document_format": "pdf",
        }

    @staticmethod
    def _process_text(data: bytes) -> tuple[str, Dict]:
        return data.decode("utf-8", errors="replace"), {"document_format": "text"}

    @staticmethod
    def _process_docx(data: bytes) -> tuple[str, Dict]:
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError("python-docx is not installed.") from exc

        document = Document(io.BytesIO(data))
        parts: List[str] = []

        for paragraph in document.paragraphs:
            value = paragraph.text.strip()
            if value:
                parts.append(value)

        for table in document.tables:
            parts.append("--- TABLE ---")
            for row in table.rows:
                parts.append(" | ".join(cell.text.strip() for cell in row.cells))

        return "\n".join(parts), {
            "document_format": "docx",
            "paragraph_count": len(document.paragraphs),
            "table_count": len(document.tables),
        }

    def _limit_text(self, text: str, result: MediaResult) -> str:
        if len(text) <= self.max_text_chars:
            return text
        result.add_warning("Document text exceeded the configured maximum and was truncated.")
        return text[: self.max_text_chars] + "\n\n[Document text truncated.]"
